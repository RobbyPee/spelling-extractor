from docx import Document
from docx.shared import Pt

def create_spelling_doc(spellings, out_path='Next_Week_Spellings.docx'):
    doc = Document()
    for word in spellings:
        run = doc.add_paragraph().add_run(word)
        run.font.size = Pt(36)
    doc.save(out_path)
    return out_path
